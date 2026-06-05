#!/usr/bin/env python3
"""Build the trial 5 / trial 6 retrospective + C3-symmetric design plan report."""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/data/binder_software/BindCraft/runs/15delt3")
ASSETS = ROOT / "report_assets"
TRIAL6 = ROOT / "trial_6_may_20"
RMSF_PNG = TRIAL6 / "rmsf_by_regions.png"
HEXAMER_PNG = TRIAL6 / "hexamer(white)_trimmer(red)_hotspot(green).png"
OUT_DOCX = ROOT / "AAV_T3_binder_design_report.docx"

doc = Document()

# ----- Document-wide style: Times New Roman 12 pt, no spacing oddities -----
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
# East-Asian variant (some Word installs need this for full font compliance)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:ascii"), "Times New Roman")
rfonts.set(qn("w:hAnsi"), "Times New Roman")
rfonts.set(qn("w:eastAsia"), "Times New Roman")
rfonts.set(qn("w:cs"), "Times New Roman")

paragraph_format = normal.paragraph_format
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

# Heading styles → Times New Roman, sensible sizing
for name, size, bold in [("Heading 1", 16, True), ("Heading 2", 14, True), ("Heading 3", 12, True)]:
    s = styles[name]
    s.font.name = "Times New Roman"
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(4)
    s_rpr = s.element.get_or_add_rPr()
    s_rfonts = s_rpr.find(qn("w:rFonts"))
    if s_rfonts is None:
        s_rfonts = OxmlElement("w:rFonts")
        s_rpr.append(s_rfonts)
    s_rfonts.set(qn("w:ascii"), "Times New Roman")
    s_rfonts.set(qn("w:hAnsi"), "Times New Roman")
    s_rfonts.set(qn("w:eastAsia"), "Times New Roman")
    s_rfonts.set(qn("w:cs"), "Times New Roman")

# Standard margins (1 inch on all sides)
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

# Helpers ------------------------------------------------------------------

def set_run_font(run, name="Times New Roman"):
    run.font.name = name
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)

def p(text="", bold=False, italic=False, align=None, style=None):
    para = doc.add_paragraph(style=style)
    if align:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run)
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def heading(text, level):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    set_run_font(run)
    # Heading font sizes
    sizes = {1: 16, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def figure(path, width_in, caption_text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption_text)
    set_run_font(cap_run)
    cap_run.font.size = Pt(11)
    cap_run.italic = True

def _shade_cell(cell, hex_color):
    """Apply a background fill to a table cell (e.g. 'D9E2F3')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def make_table(header, rows, col_widths=None):
    """Plain grid table with manually-shaded header. Avoids LibreOffice's
    quirky empty-header-row repeat when a styled header table breaks across pages."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    table.autofit = False
    # Header row — shaded light blue, bold, TNR
    for j, htext in enumerate(header):
        cell = table.rows[0].cells[j]
        _shade_cell(cell, "D9E2F3")
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(htext)
        set_run_font(run)
        run.bold = True
        run.font.size = Pt(11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body — plain, TNR
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_run_font(run)
            run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Column widths
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    return table

# =========================================================================
# TITLE
# =========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run("Design of a Hexamer-Specific Binder for the AAV T=3 Capsid:\nRetrospective on Trials 5–6 and Plan for C3-Symmetric Trimer Design")
set_run_font(trun)
trun.font.size = Pt(15)
trun.bold = True

subt = doc.add_paragraph()
subt.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subt.add_run("Internal technical report — May 2026")
set_run_font(sub_run)
sub_run.font.size = Pt(11)
sub_run.italic = True

# =========================================================================
# PART 1: TRIAL 5 AND TRIAL 6
# =========================================================================
heading("Part 1: Trial 5 and Trial 6 Retrospective", 1)

# --- 1.1 Choice of hotspot residues ------------------------------------
heading("1.1 Choice of hotspot residues (105–115)", 2)

p("The target for binder design was the AAV t3 capsid hexamer (\"rosette\"), a six-chain assembly that arises only in the T=3 icosahedral architecture and is absent from T=1 particles. The intent is therefore a binder that engages specifically this hexameric arrangement and not the dimeric asymmetric unit shared with T=1.")

p("Within each VP3 chain, the residue range 105–115 was selected as the hotspot patch on the basis of three independent observations:")

# Bulleted list (proper bullets via Word's List Bullet style)
for s in [
    "Functional annotation: the region encompasses the \"T3 hexamer knob,\" a surface loop that is structurally unique to the hexameric assembly and lies adjacent to several documented A20-antibody epitopes (Figure 1). Targeting this knob therefore engages a feature that physically exists only when the hexamer is assembled.",
    "Backbone dynamics: per-residue root-mean-square fluctuation (RMSF) from molecular dynamics indicates a moderately mobile region with a clear local maximum (~0.20 nm) within the 105–115 window relative to the global baseline (~0.08 nm). This level of flexibility is consistent with a solvent-exposed loop accessible to a binder, without being so mobile as to lack a definable target shape (Figure 1).",
    "Spatial position: a top-down view of the hexamer with the 105–115 region highlighted (Figure 2) confirms that the patch lies on the upper face of the assembly, clustered around the three-fold symmetry axis, where six copies of the loop form a near-coplanar rosette at z ≈ 180.8 Å.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

# Figure 1
figure(RMSF_PNG, 6.4,
       "Figure 1. Average backbone RMSF per residue of the VP3 monomer in the hexamer. The orange band ('T3 hexamer knob') marks residues 105–115, which display a clear local mobility peak (~0.20 nm) above the global baseline (~0.08 nm). Pink bands mark documented A20-antibody epitopes flanking the region; the blue band marks the 3-fold axis knob.")

# Figure 2
figure(HEXAMER_PNG, 5.2,
       "Figure 2. Top-down view of the AAV t3 hexamer. The full chain backbone is in white; the trimmed region used as the target is in red; the hotspot residues 105–115 across all six chains are in bright green. The green patches sit on the upper face of the assembly, clustered around the three-fold symmetry axis.")

# --- 1.2 Trial 5 ---------------------------------------------------------
heading("1.2 Trial 5 — hotspots on chains A and F", 2)

p("Trial 5 placed hotspots on chains A and F (residues 105–115 of each). These chains belong to different dimer-pairs (A↔E and C↔F), so a binder engaging both was expected to be intrinsically hexamer-specific.")

p("The configuration failed completely. After 42 trajectories (~24 GPU hours), no design reached the relaxation stage. Every trajectory ended in Trajectory/Clashing. The root cause is geometric: the contact-loss in BindCraft pulls the binder toward the centroid of the requested hotspots, which for A and F lies at (5.5, –4.0, 180.9) Å. This point is 7.2 Å from the centroid of chain C's 105–115 patch (Figure 3, dashed line) — that is, inside chain C's residue cloud. No clash-free binder placement exists for this hotspot specification.")

# --- 1.3 Trial 6 ---------------------------------------------------------
heading("1.3 Trial 6 — hotspots on chains A, C, and E", 2)

p("Trial 6 spread hotspots across three chains, one from each \"upper-class\" position (A, C, E). The binder anchor centroid lies at (8.1, 1.2, 180.8) Å, which is 17.6 Å from the closest non-hotspot patch (chain F). On purely geometric grounds the configuration should permit a clash-free fit (Figure 3, solid arrow vs blue cross).")

p("Empirically, the trajectory funnel improved markedly over Trial 5 but remains insufficient (Figure 4). The run was allowed to continue for five days on two GPUs in parallel (≈240 GPU-hours total) and was terminated on 2026-05-25. Over that window approximately 719 trajectories were attempted; 715 (99.4%) failed at the trajectory stage — 394 ended in Trajectory/Clashing and 281 were killed at the Stage-1 logits pLDDT filter. Only 4 trajectories (0.6%) completed scoring and entered the MPNN sequence-design stage. ProteinMPNN generated 80 candidate sequences (20 per surviving trajectory); every single one was rejected at AlphaFold validation, with 80/80 failing the interface pTM threshold and 80/80 failing the interface pAE threshold. No design reached the Accepted stage.")

# Figure 3 (geometry diagram)
figure(ASSETS / "fig_geometry.png", 5.0,
       "Figure 3. Hotspot patch geometry of the hexamer (top-down). Coloured circles mark the 105–115 centroid of each chain; grey segments connect dimer-pairs (A↔E, B↔D, C↔F; 6.2 Å each). The red cross marks the Trial 5 binder anchor (midpoint of A and F), only 7.2 Å from chain C's residues — a guaranteed clash. The blue cross marks the Trial 6 anchor (centroid of A, C, E); its 17.6 Å clearance from chain F permits placement but still requires the binder backbone to navigate around B, D, F when wrapping toward A, C, E.")

# Figure 4 (funnel)
figure(ASSETS / "fig_funnel.png", 6.4,
       "Figure 4. Trajectory funnel comparison after the wall-clock period indicated for each trial. Trial 5 stalled at Stage 1; Trial 6 produces a non-zero fraction of confident initial folds, but only one passes the post-hallucination relaxation, and zero reach the MPNN / acceptance stages.")

# --- 1.4 Why are they failing? -----------------------------------------
heading("1.4 Why are these designs failing", 2)

p("The combined evidence from the two trials and the per-residue SASA analysis points to a single mechanistic cause: a single-chain binder cannot simultaneously engage multiple 105–115 patches of this hexamer without colliding with the intervening chains, because the residues themselves face partially inward and the inter-patch distances exceed what a small binder can bridge cleanly.")

# Table 1 — failure mode breakdown
p("", )  # spacer
p("Table 1. Trajectory outcomes after the elapsed wall-clock period of each trial.", italic=True)

make_table(
    ["Metric", "Trial 5 (A+F)", "Trial 6 (A+C+E)"],
    [
        ["Trajectories started", "42", "≈719"],
        ["Failed Stage 1 (pLDDT too low)", "—", "318 (44%)"],
        ["Ended in Trajectory/Clashing", "42 (100%)", "394 (55%)"],
        ["Reached Trajectory scoring", "0", "4 (0.6%)"],
        ["MPNN sequences designed", "0", "80"],
        ["MPNN sequences passing AF2 validation", "0", "0 (0/80 on both i_pTM and i_pAE)"],
        ["Accepted designs", "0", "0"],
        ["GPUs used (in parallel)", "1", "2"],
        ["Wall-clock until termination", "≈24 h", "≈120 h (5 days)"],
        ["GPU-hours consumed", "≈24", "≈240"],
        ["Projected time to 50 designs", "intractable", "intractable"],
    ],
    col_widths=[2.7, 1.6, 1.6],
)
p("")

# SASA figure
figure(ASSETS / "fig_sasa.png", 6.4,
       "Figure 5. Relative solvent-accessible surface area (rSASA) of residues 105–115 in the hexamer context, separated by the two symmetry classes (A/D/F and B/C/E). The two classes have complementary burial patterns: residues 108 and 110 are buried in the A/D/F class, while residues 112–114 are buried in the B/C/E class. Across both classes 42 of 66 residues remain solvent-exposed (rSASA > 0.30), so the patch is targetable in principle.")

p("Three numerical observations summarise the failure:")
for s in [
    "Trial 5 has a hard geometric obstruction: 7.2 Å midpoint–to–chain-C distance places the requested binder anchor inside chain C, which is below typical Cα–Cα packing distances (~5 Å). No clash-free pose exists.",
    "Trial 6 has a softer geometric obstruction: the 17.6 Å midpoint clearance is safe at the anchor itself, but the binder's backbone must still reach two further patches that lie 18.6 Å apart, and the path to those patches is partially obstructed by the dimer partners (B, D, F). This is reflected in the dominant failure modes after 5 days of compute: 394 of 719 trajectories (55%) clashed irrecoverably, and of the four that did reach scoring, all 80 of the ProteinMPNN-redesigned sequences failed both the interface pTM and interface pAE thresholds at AlphaFold validation — meaning AlphaFold-multimer disagreed with BindCraft's hallucination that an interface was actually formed.",
    "The bottleneck is therefore not the choice of epitope (42 of 66 hotspot residues are solvent-exposed) or BindCraft's tuning, but the impossibility of producing a single-chain binder whose topology satisfies multi-patch contacts on a C3-symmetric target.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

p("Conclusion of Part 1: single-chain hallucination with BindCraft has been exhausted as a productive route to a hexamer-specific binder for this target. A binder whose topology is itself C3-symmetric — i.e. three identical subunits arranged around the same axis as the target — converts the multi-patch problem into three independent single-patch problems, sidesteps the bridging-distance obstruction entirely, and additionally enforces hexamer-specificity by construction. The remainder of this report develops that approach.")

# --- 1.5 Verdict on continued BindCraft use --------------------------------
heading("1.5 Verdict: should we continue using BindCraft?", 2)

p("The question is whether BindCraft should remain the engine for this project going forward. The answer is target-specific: BindCraft should be retired for the AAV t3 hexamer target, but is not being abandoned as a tool.")

heading("Arguments against continuing BindCraft on this target", 3)
for s in [
    "Empirical yield is effectively zero. After 264 GPU-hours combined across Trial 5 (24 h) and Trial 6 (240 h on two GPUs), exactly 0 designs have reached the Accepted bucket. The trajectory funnel has not just been slow — it has produced no usable output at all.",
    "The failure is mechanistic, not parametric. The 99.4% trajectory-stage failure rate in Trial 6 is dominated by hard geometric clashes (394 of 715 failed trajectories) and low confidence at Stage 1 (281). These reflect the impossibility of placing a single-chain backbone across the multi-patch C3 target geometry, not a loss-weight or hyperparameter setting that further tuning would recover.",
    "The MPNN stage independently confirms the diagnosis. The four trajectories that did survive scoring produced 80 ProteinMPNN sequences; all 80 failed AlphaFold-multimer validation on both i_pTM and i_pAE. That is, even when BindCraft's own hallucination converged, the structural prediction step disagreed with it. The designs were not real interfaces.",
    "Projected throughput remains intractable. At the observed rate (0 accepted designs in 5 days, two GPUs), reaching the requested 50 final designs would require infinite wall-clock; even a single accepted design has not been observed.",
    "BindCraft is the wrong topology engine for symmetric, multi-patch targets. The tool is designed for single-chain hallucination against a fixed receptor; it has no native concept of symmetry constraints. Forcing it onto a C3-symmetric target is using the tool against its design assumptions.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

heading("Arguments for keeping BindCraft in the toolbox", 3)
for s in [
    "BindCraft remains a sound choice for single-chain, single-patch targets. The prior PDL1 work in this repository (settings_target/PDL1.json) and the broader BindCraft literature both demonstrate productive yields on non-symmetric targets. Nothing in Trials 5–6 contradicts that.",
    "The infrastructure investment is reusable. Filter sets, advanced-settings JSONs, the reporting pipeline (rank_current.py, bindcraft_report.py), and the per-trial directory convention can all be applied to any future target without rework.",
    "The trajectory and failure data collected here are useful as a baseline. Future symmetric-target attempts (whether with RFdiffusion or any other engine) can be benchmarked against the 0.6% scoring-stage yield observed here to confirm they are actually doing better, rather than merely seeming to.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

heading("Recommendation", 3)
p("Stop running BindCraft against the AAV t3 hexamer. Trial 6 was terminated on 2026-05-25 after 5 days of compute on the basis of the data summarised above. Future BindCraft work on this machine should be reserved for non-symmetric targets where its single-chain hallucination matches the problem topology. For the hexamer-specific binder, proceed directly to the C3-symmetric pipeline described in Part 2, which uses RFdiffusion for backbone generation, an explicit geometric rotation for symmetry enforcement, ProteinMPNN with tied positions for sequence design, and AlphaFold-multimer for validation — none of which require or use BindCraft.")

doc.add_page_break()

# =========================================================================
# PART 2: C3-SYMMETRIC DESIGN STRATEGY
# =========================================================================
heading("Part 2: C3-Symmetric Trimer Binder — Design Strategy", 1)

# --- 2.1 Big picture ----------------------------------------------------
heading("2.1 The biological objective", 2)

p("The AAV capsid can be viewed as a soccer-ball-like assembly built from VP3 protein subunits. Two icosahedral arrangements are relevant here:")
for s in [
    "T=1 capsid: VP3 subunits arrange as 30 dimers. Dimer-only contacts dominate.",
    "T=3 capsid: VP3 subunits arrange as three dimers clustered around each three-fold axis, forming the \"rosette\" of six chains that is the structural target of this design. This rosette does not exist in T=1.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

p("The design goal is therefore a binder that engages only the six-chain rosette and not any individual dimer pair. A binder that also recognises an isolated dimer would adhere to T=1 capsids as well and fail the specificity requirement.")

# --- 2.2 Why single-chain is hard --------------------------------------
heading("2.2 Why a single-chain binder is the wrong topology", 2)

p("As established in Part 1, a single-chain binder asked to contact multiple 105–115 patches across the hexamer collides with the intervening chains. Reducing the number of contacted patches makes the design tractable but eliminates the cross-dimer-pair contacts that distinguish the hexamer from a dimer. Conversely, contacting the cross-dimer-pair patches requires the binder backbone to thread through space already occupied by other chains.")

p("The contact-loss function in BindCraft is an OR-style sum over hotspot positions: a binder is rewarded for contacting any subset of hotspots and is not required to contact all of them. Consequently, even when a single-chain binder is large enough to span the hexamer geometrically, the optimiser tends to collapse it into one chain's pocket and ignore the others — producing a design that, however good its computed metrics, would not in fact be hexamer-specific in vitro.")

# --- 2.3 The C3 idea ---------------------------------------------------
heading("2.3 The C3-symmetric trimer solution", 2)

p("The hexamer itself has exact three-fold rotational symmetry: rotation by 120° about the axis through (0, 0, 180.8) Å maps each dimer-pair onto the next (verified to within 0.3° using the dimer-pair centroids). If the binder is constructed to share this C3 symmetry — three identical subunits arranged 120° apart around the same axis — then each subunit can be a small, geometrically simple binder against one chain, while the assembled trimer engages all three target chains simultaneously.")

p("This converts the design problem into three coupled tasks:")
for s in [
    "Design one binder subunit that engages one chain of the hexamer cleanly. This is a standard single-target, single-chain design problem with no bridging requirement.",
    "Replicate this subunit by 120° and 240° rotation about the verified C3 axis. This is a deterministic geometric operation; no machine learning is involved.",
    "Verify, via AlphaFold-multimer prediction, that the assembled trimer folds as designed when bound to the full hexamer; and crucially, that it does not bind when the target is reduced to a single dimer pair.",
]:
    para = doc.add_paragraph(style="List Bullet Number") if False else doc.add_paragraph(style="List Number")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

p("On the full hexamer all three subunits engage their partners simultaneously, producing a multivalent interaction with combined affinity well below the dissociation constant of any one subunit alone. On an isolated dimer at most one subunit finds a partner — the other two are unmatched — and the multivalent advantage collapses, producing the desired loss of binding.")

# --- 2.4 RFdiffusion limitation -------------------------------------
heading("2.4 Why we will not use RFdiffusion's symmetric PPI mode directly", 2)

p("RFdiffusion supports symmetric design natively (C3, C4, dihedral, tetrahedral, octahedral, icosahedral) and supports hotspot-biased PPI design natively. However, combining the two — symmetric protein-protein interaction design with hotspot bias — is explicitly flagged in the official RFdiffusion documentation as unreliable in the current release. The maintainers write that hotspot residues \"seem to interact weirdly with the symmetric PPI potentials\" and that improved potentials are planned but not released. In practice, designs from this combined mode show low binder-target interface confidence and frequent failure to satisfy the symmetry constraint at the same time as the hotspot constraint.")

p("To avoid this software limitation, the strategy decomposes the symmetric problem into a non-symmetric design step followed by an explicit geometric replication step:")
for s in [
    "Phase 1: design a single binder subunit against one chain of the hexamer using regular (non-symmetric) PPI mode. Hotspots work reliably here.",
    "Phase 2: rotate the subunit by 120° and 240° about the C3 axis to produce the second and third subunits. Fuse the three by flexible peptide linkers into a single polypeptide chain.",
    "Phase 3: ProteinMPNN with tied positions across the three subunits to assign a single sequence, preserving symmetry at the sequence level.",
    "Phase 4: AlphaFold-multimer prediction against (i) the full hexamer, (ii) an isolated dimer pair — the latter being the discriminating specificity test.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

p("The end-to-end workflow, including expected per-phase yields, key parameters, and validation filters, is summarised in Figure 6.")

figure(ASSETS / "fig_workflow.png", 6.4,
       "Figure 6. Proposed workflow for C3-symmetric trimer binder design. The pipeline is decomposed into four phases plus a terminal specificity test. Phase 1 generates one binder subunit against chain A by standard (non-symmetric) RFdiffusion PPI; Phase 2a replicates the subunit by deterministic 120° / 240° rotation around the verified hexamer C3 axis and fuses the three copies into a single polypeptide via flexible (GGGGS)ₙ linkers; Phase 2b validates the assembled trimer with AlphaFold-multimer against the full hexamer and applies four filters. Phase 3 designs sequences with ProteinMPNN under tied-position constraints to preserve sequence-level symmetry. Phase 4 re-validates each MPNN sequence with the same filters. The terminal acid test re-predicts each survivor against an isolated dimer pair (chains A + E only); a confirmed hexamer-specific design must lose at least 0.15 in interface pTM and have at least one subunit with negligible interface SASA when the hexamer is reduced to a dimer. Expected yield (n) at each stage is annotated on the right; total wall-clock for the pilot run is approximately 6–8 hours on a single GPU.")

# --- 2.5 The five design decisions in detail ---------------------------
heading("2.5 Detailed design choices", 2)

heading("(a) Two-phase strategy in place of direct symmetric design", 3)
p("Phase 1 uses RFdiffusion's well-validated single-chain PPI mode to design one subunit against chain A of the hexamer, with the surrounding chains (B–F) included in the input as fixed steric context but not as hotspot targets. This avoids the buggy symmetric-PPI mode of RFdiffusion entirely. The C3 replication that produces subunits 2 and 3 is performed in Phase 2 by a deterministic rotation of the Phase 1 backbone about the verified hexamer C3 axis. No machine learning is involved in the symmetry step, which removes the principal source of failure in the alternative approach.")

heading("(b) Subunit length 60–90 residues, six exposed hotspot residues on chain A", 3)
p("Each subunit needs to engage one chain at one local patch (residues 105–115 of chain A), so it can be small. The chosen length range, 60–90 residues, is comparable to a small de novo mini-binder (for comparison: GFP is 239 residues, an antibody scFv is ~250 residues). The hotspot list is restricted to the six residues that are solvent-exposed in the hexamer context (rSASA > 0.30): A105, A107, A109, A111, A114, A115. The four partially or fully buried residues (A106, A108, A110, A113) are excluded because they cannot be reached without occluding the dimer-partner interface. Restricting hotspots to genuinely accessible residues sharpens the loss signal and increases the per-trajectory success rate.")

heading("(c) Geometric replication around the verified C3 axis at (0, 0, 180.8) Å", 3)
p("Pairwise centroid analysis of the six 105–115 patches confirms that they lie at three positions exactly 120° apart around the vertical axis through (0, 0, 180.8) Å, with radial deviations below 0.3°. The Phase 2 rotation can therefore be performed about the canonical z axis after recentering the binder subunit on this point — a trivial linear transformation. Each of the three replicated subunits lands on a different chain of the same symmetry class (A → C → E, or by symmetry B → D → F), so the resulting trimer engages one chain from each dimer-pair of the hexamer, which is the condition for hexamer-specificity.")

heading("(d) Tied positions in ProteinMPNN to preserve symmetry at sequence level", 3)
p("After RFdiffusion produces a backbone and the C3 replication produces three copies of it, the three copies must remain identical not only at the backbone level but also at the sequence level — otherwise the assembly is no longer a true homotrimer. ProteinMPNN by default treats the three subunits as three independent sequence-design problems, which would assign different residues at corresponding positions and break the symmetry. The tied_positions option constrains residue i of subunit 1 to equal residue i of subunits 2 and 3 for all i, so ProteinMPNN designs a single sequence that is applied to all three copies. The same tying convention is already in use in the trial_2B pipeline on this machine for an analogous dimer-of-dimers redesign.")

heading("(e) Dimer-only specificity test as the final validation criterion", 3)
p("After all other filters (binder pLDDT > 0.70, interface pTM > 0.65, per-subunit interface SASA > 200 Å², backbone RMSD vs design < 3.0 Å), each surviving design is re-predicted by AlphaFold-multimer against a target consisting of only a single dimer pair (chains A and E) rather than the full hexamer. A truly hexamer-specific binder must show a substantial drop in interface confidence under this condition — interface pTM should decrease by at least 0.15, reflecting that two of the three subunits have lost their partners. Designs that retain comparable interface pTM against the dimer are classified as accidentally promiscuous and discarded, regardless of how well they scored on the hexamer. This test is the operational definition of hexamer-specificity used by this project.")

# --- 2.6 Homotrimer assembly --------------------------------------------
heading("2.6 Producing the homotrimer experimentally", 2)

p("The C3-symmetric design produced in silico must be expressed and assembled as an actual trimer in solution. Four established strategies are available; the recommended choice for this project is single-chain fusion, with foldon trimerisation as a documented fallback.")

# Table 2 — homotrimer strategies
p("Table 2. Experimental routes to a homotrimeric binder, ranked by ease of implementation.", italic=True)
make_table(
    ["Strategy", "Mechanism", "Pros", "Cons", "Use case"],
    [
        ["Single-chain genetic fusion",
         "Three subunits joined head-to-tail by flexible (GGGGS)n linkers, expressed as one polypeptide",
         "One gene, one protein, guaranteed 1:1:1 stoichiometry",
         "Slightly larger expressed protein (~226 aa)",
         "Default for this project"],
        ["Foldon trimerisation tag",
         "Each subunit fused to a 27-aa T4 fibritin domain that self-trimerises in solution",
         "Smaller individual chains; tag well-characterised",
         "Adds non-functional protein domain; may be immunogenic for therapeutic use",
         "Fallback if single-chain fusion fails to fold"],
        ["Engineered disulfide bridges",
         "Designed cysteine pairs at adjacent-subunit interfaces form covalent S–S bonds",
         "Covalently locked, very stable",
         "Requires cysteine-free starting subunits and tight geometric design",
         "Future stabilisation step"],
        ["Self-assembling coiled-coil interface",
         "Inter-subunit interface designed as a known trimeric coiled-coil",
         "Clean, no extra protein domain",
         "Most demanding computational design step",
         "Optional later optimisation"],
    ],
    col_widths=[1.7, 2.0, 1.5, 1.6, 1.2],
)
p("")

p("Single-chain fusion is the recommended starting point. Three Phase-1 subunits, each ~70 residues, joined by two (GGGGS)x linkers of approximately 8 residues, produces a single polypeptide of ~226 residues. The protein folds intramolecularly: the first subunit folds, the linker bends, the second subunit folds in position, and so on. Because the linker geometry is constrained by the spatial positions required for C3 alignment, the protein cannot adopt an off-target topology without paying a folding-energy penalty. AlphaFold-multimer can predict and validate the fusion construct in one shot.")

p("Standard production workflow for the single-chain fusion is summarised in Table 3.")

# Table 3 — production workflow
p("Table 3. Wet-lab production workflow for the single-chain fusion binder.", italic=True)
make_table(
    ["Step", "Action", "Expected output / metric"],
    [
        ["1. Design", "Run Phase 1–4 of the pre-binder pipeline; output a single-chain fusion PDB",
         "Validated design with all four AF2 filters passed and dimer-only specificity confirmed"],
        ["2. Gene synthesis", "Synthesise codon-optimised gene (~700 bp) at Twist or IDT; include N-terminal His6 tag and stop codon",
         "1 week turnaround; cost ≈ USD 100–200"],
        ["3. Cloning", "Insert into pET28a (E. coli) or pcDNA3 (mammalian) by standard restriction cloning",
         "Sequenced clone within ~3 days"],
        ["4. Expression", "Transform BL21(DE3); induce with IPTG; harvest after 4–16 h at 18–37 °C",
         "5–50 mg / L culture (typical mini-binder yield)"],
        ["5. Purification", "IMAC nickel column followed by size-exclusion chromatography",
         "Single SEC peak at expected MW (~25 kDa)"],
        ["6. Quality control", "SDS-PAGE, electrospray mass spectrometry, SEC-MALS",
         "Single band, exact mass within 1 Da, MALS confirms monomer of the fusion (= one trimer-shaped molecule)"],
        ["7. Functional assay", "Surface plasmon resonance (SPR) or biolayer interferometry (BLI) against (a) the full hexamer, (b) an isolated dimer",
         "Hexamer binding KD ≤ 100 nM; dimer KD ≥ 10× weaker — the experimental acid test of hexamer-specificity"],
    ],
    col_widths=[1.1, 3.0, 2.9],
)
p("")

p("Total time from validated design to functional protein is approximately 6–8 weeks, at a reagent cost of approximately USD 500–2000 depending on existing infrastructure. Because the design phase produces a regular single-chain protein — albeit somewhat longer than usual — no specialised expression or assembly machinery is needed.")

# --- 2.7 Decision criteria ---------------------------------------------
heading("2.7 Decision criteria for execution", 2)

p("The C3-symmetric pipeline is staged under the pre-binder folder and is not yet executed. Activation criteria are:")
for s in [
    "Trial 6 yields fewer than three accepted designs within the next 24 hours of compute (current projected throughput of ~30 days to 50 designs makes this very likely).",
    "Trial 6 yields designs that fail the dimer-only specificity test (i.e., bind the dimer comparably to the hexamer), regardless of other metrics.",
    "An independent need arises for a hexamer-specific binder with provable mechanism, in which case the C3-symmetric design is preferred because the symmetry argument is constructive rather than empirical.",
]:
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(s)
    set_run_font(run)
    run.font.size = Pt(12)

p("If trial 6 produces ≥3 accepted designs that also pass the dimer specificity test, the C3-symmetric pipeline remains as a parallel approach worth pursuing for mechanistic rigor, but is not strictly required.")

# --- 2.8 Summary --------------------------------------------------------
heading("2.8 Summary", 2)

p("Trials 5 and 6 demonstrated experimentally that single-chain binder design with BindCraft does not yield a hexamer-specific binder for the AAV t3 capsid hexamer within tractable compute. The constraint is geometric: a single binder cannot simultaneously engage multiple 105–115 patches on a C3-symmetric target without colliding with intervening chains.")

p("A C3-symmetric trimer binder, in which the binder's topology matches the symmetry of the target, addresses both obstacles. Three identical subunits, each designed against a single chain by standard methods, are replicated by exact 120° rotation about the verified hexamer C3 axis and fused into a single polypeptide by flexible linkers. The construct is a regular single-chain protein at the wet-lab interface and a multivalent C3-symmetric binder at the structural level. Specificity for the hexamer over an isolated dimer arises naturally from the loss of two of three subunit interactions when the target's symmetry is removed.")

p("The pipeline is fully staged at /data/binder_software/BindCraft/runs/15delt3/pre-binder/, with all scripts, configurations, and documentation prepared. Execution awaits confirmation that the BindCraft route has been exhausted or, equivalently, the elapse of the next ~24 hours of Trial 6 without an inflection in its success rate.")

# Save
doc.save(str(OUT_DOCX))
print(f"Saved to {OUT_DOCX}")
print(f"Size: {OUT_DOCX.stat().st_size / 1024:.1f} kB")
